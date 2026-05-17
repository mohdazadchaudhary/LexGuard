import asyncio
import io
import pypdf
import docx
from app.ai_orchestration.gemini_client import get_gemini_client, GeminiClient


class AnalysisService:
    """
    Orchestrates document parsing and AI-powered legal analysis.

    Follows Single Responsibility: parsing is delegated to extract_* methods,
    AI calls are delegated to the GeminiClient.
    """

    def __init__(self) -> None:
        self.ai_client: GeminiClient = get_gemini_client()

    # ------------------------------------------------------------------
    # Document extraction helpers
    # ------------------------------------------------------------------

    def extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extracts plain text from a PDF file.

        @param pdf_bytes - Raw bytes of the PDF file.
        @returns Concatenated text from all pages.
        @throws ValueError if the PDF cannot be parsed.
        """
        text = ""
        try:
            reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text.strip()
        except Exception as exc:
            raise ValueError(f"Failed to parse PDF: {exc}") from exc

    def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """
        Extracts plain text from a DOCX file, including table cells.

        @param docx_bytes - Raw bytes of the DOCX file.
        @returns Extracted text as a newline-separated string.
        @throws ValueError if the DOCX cannot be parsed.
        """
        try:
            doc = docx.Document(io.BytesIO(docx_bytes))
            full_text: list[str] = []
            for para in doc.paragraphs:
                if para.text:
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            return "\n".join(full_text).strip()
        except Exception as exc:
            raise ValueError(f"Failed to parse DOCX: {exc}") from exc

    def extract_text_from_txt(self, txt_bytes: bytes) -> str:
        """
        Decodes plain text file bytes, trying UTF-8 then Latin-1 as fallback.

        @param txt_bytes - Raw bytes of the text file.
        @returns Decoded string content.
        """
        try:
            return txt_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            return txt_bytes.decode("latin-1").strip()

    async def extract_text_from_image_async(self, image_bytes: bytes) -> str:
        """
        Delegates image OCR to the GeminiClient's multimodal model.

        @param image_bytes - Raw bytes of the image (PNG, JPG, JPEG).
        @returns Extracted text from the image.
        @throws ValueError if OCR fails.
        """
        return await self.ai_client.extract_text_from_image_async(image_bytes)

    # ------------------------------------------------------------------
    # Core analysis orchestration
    # ------------------------------------------------------------------

    async def process_document_text(self, text: str) -> str:
        """
        Orchestrates parallel multi-agent analysis of a document.

        Runs three specialised agents concurrently (financial, liability, privacy)
        and assembles their outputs into a unified markdown report.

        @param text - The full document text to analyse.
        @returns Formatted multi-agent intelligence report as a markdown string.
        @throws ValueError if the text is empty.
        """
        if not text or not text.strip():
            raise ValueError("Document text cannot be empty.")

        # O(1) concurrent calls — no sequential blocking
        financial_task = self.ai_client.agent_financial_async(text)
        liability_task = self.ai_client.agent_liability_async(text)
        privacy_task = self.ai_client.agent_privacy_async(text)

        results = await asyncio.gather(financial_task, liability_task, privacy_task)

        report = f"""# LexGuard Multi-Agent Intelligence Report

## 💰 Financial Risk Analysis
{results[0]}

---

## ⚖️ Legal & Liability Analysis
{results[1]}

---

## 🔒 Privacy & Data Security Analysis
{results[2]}
"""
        return report

    async def chat_about_document(self, text: str, question: str) -> str:
        """
        Answers a user question grounded in the provided document text.

        @param text - The full document text serving as the knowledge base.
        @param question - The user's natural-language question.
        @returns A concise, contract-grounded answer string.
        @throws ValueError if text or question are empty.
        """
        if not text or not text.strip():
            raise ValueError("Document text cannot be empty.")
        if not question or not question.strip():
            raise ValueError("Question cannot be empty.")

        return await self.ai_client.chat_with_contract_async(text, question)
